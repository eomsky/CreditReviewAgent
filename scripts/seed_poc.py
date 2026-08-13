"""Create reproducible POC data and representative credit-review documents."""

from __future__ import annotations

import argparse
import uuid
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from app.core.config import PROJECT_ROOT
from app.database.poc_store import index_document, initialize_database
from app.services.documents import extract_text


SAMPLE_DIR = PROJECT_ROOT / "data" / "sample_documents"
DOCUMENTS = {
    "여신심사_운영지침": [
        "신규 기업여신은 최근 3개년 재무제표와 최근 월 매출 및 결제예정액을 확인한다.",
        "부채비율 250% 이상 또는 이자보상배율 1.0배 미만 기업은 고위험 검토 대상으로 분류한다.",
        "30일 이상 연체 이력이 있으면 차주의 상환재원과 담보가치를 별도로 검증한다.",
        "심사결정에는 데이터 기준일, 사용 근거, 확인하지 못한 가정을 명시한다.",
    ],
    "업종별_리스크_가이드": [
        "건설업은 미청구공사, 우발채무, 프로젝트 파이낸싱 보증을 중점 확인한다.",
        "도소매업은 재고회전일수와 매출채권 회수기간을 확인한다.",
        "IT서비스업은 고객 집중도, 반복매출 비중, 핵심 인력 의존도를 검토한다.",
        "부동산업은 분양률, 담보인정비율, 금리 상승 민감도를 스트레스 테스트한다.",
    ],
    "담보평가_체크리스트": [
        "부동산 담보는 최근 감정가와 선순위 권리, 임대차보증금을 차감해 유효담보가를 산출한다.",
        "매출채권 담보는 채무자의 신용도, 양도 가능성 및 과거 회수율을 확인한다.",
        "기계설비는 처분 가능성, 경제적 내용연수, 이전비용을 반영한다.",
    ],
}


def _register_pdf_font() -> str | None:
    candidates = (
        Path("C:/Windows/Fonts/malgun.ttf"),
        Path("/usr/share/fonts/truetype/nanum/NanumGothic.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    )
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("PocUnicode", str(candidate)))
            return "PocUnicode"
    return None


def create_pdf(path: Path, title: str, paragraphs: list[str]) -> None:
    font = _register_pdf_font()
    if font is None:
        title = "Credit Review Operating Policy"
        paragraphs = [
            "Review three years of financial statements and recent sales.",
            "Debt ratio above 250 percent or interest coverage below 1.0 is high risk.",
            "Delinquency of 30 days or more requires repayment-source validation.",
            "Record the data date, evidence, and any unverified assumptions.",
        ]
        font = "Helvetica"
    pdf = canvas.Canvas(str(path))
    pdf.setFont(font, 16)
    pdf.drawString(60, 790, title)
    pdf.setFont(font, 10)
    y = 750
    for index, paragraph in enumerate(paragraphs, 1):
        pdf.drawString(65, y, f"{index}. {paragraph}")
        y -= 42
    pdf.save()


def create_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    document = Document()
    document.add_heading(title, 0)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph, style="List Bullet")
    document.save(path)


def create_xlsx(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "심사기준"
    sheet.append(["지표", "정상", "주의", "고위험", "조치"])
    sheet.append(["부채비율", "150% 미만", "150~250%", "250% 이상", "상환재원 추가 검증"])
    sheet.append(["이자보상배율", "2.0배 이상", "1.0~2.0배", "1.0배 미만", "금리 스트레스 테스트"])
    sheet.append(["연체일수", "0일", "1~29일", "30일 이상", "신규 취급 제한 검토"])
    sectors = workbook.create_sheet("업종한도")
    sectors.append(["업종", "권고 한도 비중", "중점 검토사항"])
    sectors.append(["건설", "15%", "PF보증·미청구공사"])
    sectors.append(["부동산", "12%", "분양률·LTV"])
    sectors.append(["IT서비스", "20%", "반복매출·고객집중도"])
    workbook.save(path)


def seed_documents() -> int:
    SAMPLE_DIR.mkdir(parents=True, exist_ok=True)
    created: list[Path] = []
    for index, (title, paragraphs) in enumerate(DOCUMENTS.items()):
        path = SAMPLE_DIR / (f"{title}.pdf" if index == 0 else f"{title}.docx")
        create_pdf(path, title, paragraphs) if index == 0 else create_docx(path, title, paragraphs)
        created.append(path)
    xlsx_path = SAMPLE_DIR / "여신정책_정량기준.xlsx"
    create_xlsx(xlsx_path)
    created.append(xlsx_path)

    for path in created:
        index_document(
            uuid.uuid5(uuid.NAMESPACE_URL, str(path)).hex,
            path.stem,
            path,
            "application/octet-stream",
            extract_text(path),
        )
    return len(created)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--force", action="store_true", help="Reserved for future destructive reseeding")
    parser.parse_args()
    initialize_database(seed=True)
    print(f"POC seed complete: {seed_documents()} documents")


if __name__ == "__main__":
    main()
