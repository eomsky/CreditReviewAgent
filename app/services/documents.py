"""Document persistence, text extraction, and indexing helpers."""

from __future__ import annotations

import io
import mimetypes
import re
import uuid
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from app.core.config import settings
from app.database.poc_store import index_document, save_uploaded_file, update_uploaded_file
from app.rag.indexer import VectorIndexer


def safe_filename(filename: str) -> str:
    return re.sub(r"[^0-9A-Za-z가-힣._-]", "_", Path(filename).name)


def _normalized_entity(value: str) -> str:
    normalized = re.sub(r"주식회사|㈜|\(주\)", "", value, flags=re.IGNORECASE)
    return re.sub(r"[^0-9A-Za-z가-힣]", "", normalized).lower()


def is_company_relevant(text: str, filename: str, company_name: str | None) -> bool:
    """Fail closed: case uploads are evidence only when the target company is named."""
    target = _normalized_entity(company_name or "")
    if len(target) < 2:
        return False
    haystack = _normalized_entity(f"{filename}\n{text}")
    return target in haystack


def persist_and_index(
    conversation_id: str,
    filename: str,
    mime_type: str,
    raw: bytes,
    case_id: str | None = None,
    *,
    target_company_name: str | None = None,
    extracted_text: str | None = None,
) -> tuple[Path, str, str]:
    stored = settings.UPLOAD_DIR / f"{uuid.uuid4().hex}_{safe_filename(filename)}"
    stored.write_bytes(raw)
    file_id = save_uploaded_file(
        conversation_id, filename, stored, mime_type, len(raw), "", case_id=case_id, status="UPLOADED"
    )
    try:
        update_uploaded_file(file_id, status="PARSING")
        text = extracted_text if extracted_text is not None else extract_text(stored, mime_type, raw)
        update_uploaded_file(file_id, status="CHUNKING", extracted_text=text)
        if target_company_name and not is_company_relevant(text, filename, target_company_name):
            update_uploaded_file(
                file_id,
                status="EXCLUDED",
                extracted_text=text,
                error_message=f"심사대상 기업({target_company_name}) 정보가 확인되지 않아 근거에서 제외했습니다.",
            )
            return stored, text, file_id
        if text.strip():
            index_document(file_id, filename, stored, mime_type, text, case_id=case_id)
            update_uploaded_file(file_id, status="EMBEDDING")
            VectorIndexer().index(file_id, filename, text, case_id)
        update_uploaded_file(file_id, status="READY")
        return stored, text, file_id
    except Exception as exc:
        update_uploaded_file(file_id, status="FAILED", error_message=type(exc).__name__)
        raise


def extract_text(path: Path, mime_type: str | None = None, raw: bytes | None = None) -> str:
    raw = raw if raw is not None else path.read_bytes()
    suffix = path.suffix.lower()
    if (mime_type or "").startswith("image/"):
        return ""
    if mime_type == "application/pdf" or suffix == ".pdf":
        return "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(raw)).pages)
    if suffix == ".docx":
        document = Document(io.BytesIO(raw))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            parts.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(parts)
    if suffix in {".xlsx", ".xlsm"}:
        workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in workbook.worksheets:
            lines.append(f"[시트: {sheet.title}]")
            lines.extend("\t".join("" if value is None else str(value) for value in row) for row in sheet.iter_rows(values_only=True))
        return "\n".join(lines)
    if suffix in {".txt", ".csv", ".md", ".json"} or (mime_type or "").startswith("text/"):
        return raw.decode("utf-8", errors="replace")
    guessed, _ = mimetypes.guess_type(path.name)
    raise ValueError(f"지원하지 않는 문서 형식입니다: {guessed or suffix}")
